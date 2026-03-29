"""DOCX renderer using python-docx.

Produces professional Word documents with:
- Branding header/footer on every page
- Styled tables with alternating rows
- Proper heading hierarchy
- Metadata block
- Page numbers

Requires the optional 'docx' dependency: pip install forgedoc[docx]
"""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from ..core import Document, Section, TableDef
    from ..styles.default import StyleConfig


def render_docx(doc: Document, style=None, **kwargs) -> bytes:
    """Render a Document to DOCX bytes."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX output. "
            "Install it with: pip install forgedoc[docx]"
        )

    docx = DocxDocument()

    # Page setup for landscape if needed
    section = docx.sections[0]
    if "landscape" in (doc.page_config.size or "").lower():
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    # Branding colors
    primary_rgb = _hex_to_rgb(style.primary_color if style else "#1a1a2e")
    accent_rgb = _hex_to_rgb(style.accent_color if style else "#16213e")
    muted_rgb = _hex_to_rgb(style.muted_color if style else "#666666")

    # Header with branding
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    if doc.branding.company_name:
        run = header_para.add_run(doc.branding.company_name)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*muted_rgb)
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_text = doc.branding.footer_text or doc.title
    run = footer_para.add_run(f"{footer_text}    ")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(*muted_rgb)
    # Add page number field
    _add_page_number(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title_para = docx.add_heading(doc.title, level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(*primary_rgb)

    if doc.subtitle:
        sub = docx.add_paragraph()
        run = sub.add_run(doc.subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*muted_rgb)

    if doc.branding.company_name and not header.paragraphs[0].text:
        p = docx.add_paragraph()
        run = p.add_run(doc.branding.company_name)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*muted_rgb)

    # Metadata table
    if doc.metadata:
        meta_items = [(k, v) for k, v in doc.metadata.items() if v]
        if meta_items:
            table = docx.add_table(rows=len(meta_items), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, (k, v) in enumerate(meta_items):
                row = table.rows[i]
                # Bold key
                key_para = row.cells[0].paragraphs[0]
                key_run = key_para.add_run(str(k))
                key_run.bold = True
                key_run.font.size = Pt(9)
                # Value
                val_para = row.cells[1].paragraphs[0]
                val_run = val_para.add_run(str(v))
                val_run.font.size = Pt(9)
            # Minimal borders on metadata table
            _set_table_borders(table, color="CCCCCC", size=4)
            docx.add_paragraph()  # Spacer

    # Sections
    for sec in doc.sections:
        _render_section_docx(docx, sec, primary_rgb, accent_rgb)

    buf = BytesIO()
    docx.save(buf)
    return buf.getvalue()


def _render_section_docx(docx, section: Section, primary_rgb, accent_rgb):
    """Render a section into the DOCX document."""
    from docx.shared import Pt, RGBColor

    if section.title:
        level = min(section.level, 4)
        heading = docx.add_heading(section.title, level=level)
        color = primary_rgb if level == 1 else accent_rgb
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)

    if section.content:
        _render_markdown_docx(docx, section.content)

    for table_def in section.tables:
        _render_table_docx(docx, table_def)

    for sub in section.subsections:
        _render_section_docx(docx, sub, primary_rgb, accent_rgb)


def _render_markdown_docx(docx, content: str):
    """Render markdown content as DOCX paragraphs with inline formatting."""
    import re as _re

    from docx.shared import Pt

    for block in _re.split(r"\n{2,}", content.strip()):
        lines = block.strip().split("\n")
        if not lines or not lines[0].strip():
            continue

        # Bullet list
        if all(_re.match(r"^[\-\*]\s+", l.strip()) for l in lines if l.strip()):
            for line in lines:
                text = _re.sub(r"^[\-\*]\s+", "", line.strip())
                p = docx.add_paragraph(style="List Bullet")
                _add_md_runs(p, text, Pt(9))
            continue

        # Numbered list
        if all(_re.match(r"^\d+\.\s+", l.strip()) for l in lines if l.strip()):
            for line in lines:
                text = _re.sub(r"^\d+\.\s+", "", line.strip())
                p = docx.add_paragraph(style="List Number")
                _add_md_runs(p, text, Pt(9))
            continue

        # Regular paragraph
        joined = " ".join(l.strip() for l in lines if l.strip())
        p = docx.add_paragraph()
        _add_md_runs(p, joined, Pt(9))


def _add_md_runs(paragraph, text: str, font_size):
    """Add runs to a paragraph with bold/italic/code formatting."""
    import re as _re

    from docx.shared import Pt, RGBColor

    # Split on markdown patterns, keeping delimiters
    # Pattern matches: **bold**, *italic*, `code`, [text](url)
    parts = _re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))", text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = font_size
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = font_size
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)
        elif part.startswith("["):
            match = _re.match(r"\[(.+?)\]\((.+?)\)", part)
            if match:
                run = paragraph.add_run(match.group(1))
                run.font.size = font_size
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 180)
            else:
                run = paragraph.add_run(part)
                run.font.size = font_size
        else:
            run = paragraph.add_run(part)
            run.font.size = font_size


def _render_table_docx(docx, table_def: TableDef):
    """Render a table with styled header and alternating rows."""
    from docx.shared import Pt, RGBColor

    ncols = len(table_def.headers)
    table = docx.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(table_def.headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Dark background on header
        _set_cell_bg(cell, "1a1a2e")

    # Data rows
    for row_idx, row_data in enumerate(table_def.rows):
        row = table.add_row()
        for i, cell_val in enumerate(row_data):
            if i < ncols:
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_val))
                run.font.size = Pt(8)
                # Alternating row background
                if row_idx % 2 == 1:
                    _set_cell_bg(cell, "fafafa")

    if table_def.caption:
        p = docx.add_paragraph()
        run = p.add_run(table_def.caption)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(102, 102, 102)
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Convert hex color string to (R, G, B) tuple."""
    h = hex_color.lstrip("#")
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    """Set cell background color via XML."""
    from docx.oxml.ns import qn

    shading = cell._element.get_or_add_tcPr().makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): hex_color,
            qn("w:val"): "clear",
        },
    )
    cell._element.get_or_add_tcPr().append(shading)


def _set_table_borders(table, color="auto", size=4):
    """Set table borders via XML."""
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "single",
                qn("w:sz"): str(size),
                qn("w:space"): "0",
                qn("w:color"): color,
            },
        )
        borders.append(element)
    tblPr.append(borders)


def _add_page_number(paragraph):
    """Add a page number field to a paragraph."""
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = run2._element.makeelement(qn("w:instrText"), {})
    instrText.text = " PAGE "
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fldChar2)
